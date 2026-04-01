from flask import Blueprint, request, jsonify
from utils.aws_client import create_aws_client
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import base64
import io
from PIL import Image, ImageDraw, ImageFont
import requests

cloudformation_bp = Blueprint('cloudformation', __name__)

def create_aws_logo_image():
    """Create a simple AWS logo-style image with text"""
    try:
        # Create a 800x400 image with white background
        img = Image.new('RGB', (800, 400), color='white')
        draw = ImageDraw.Draw(img)
        
        # AWS orange color
        aws_orange = '#FF9900'
        text_color = '#232F3E'  # AWS dark blue
        
        # Draw AWS-style rectangle
        draw.rectangle([50, 50, 750, 350], outline=aws_orange, width=3)
        
        # Try to use a default font, fallback to basic if not available
        try:
            font_large = ImageFont.truetype('/System/Library/Fonts/Arial.ttf', 36)
            font_medium = ImageFont.truetype('/System/Library/Fonts/Arial.ttf', 24)
        except:
            try:
                font_large = ImageFont.load_default()
                font_medium = ImageFont.load_default()
            except:
                font_large = None
                font_medium = None
        
        # Add AWS text
        aws_text = "AWS"
        if font_large:
            bbox = draw.textbbox((0, 0), aws_text, font=font_large)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
        else:
            text_width, text_height = 60, 30
        
        aws_x = (800 - text_width) // 2
        aws_y = 120
        draw.text((aws_x, aws_y), aws_text, fill=aws_orange, font=font_large)
        
        # Add main message
        main_text = "Infrastructure As Code Change:"
        sub_text = "No test plan available"
        
        if font_medium:
            # Main text
            bbox = draw.textbbox((0, 0), main_text, font=font_medium)
            main_width = bbox[2] - bbox[0]
            main_x = (800 - main_width) // 2
            draw.text((main_x, 200), main_text, fill=text_color, font=font_medium)
            
            # Sub text
            bbox = draw.textbbox((0, 0), sub_text, font=font_medium)
            sub_width = bbox[2] - bbox[0]
            sub_x = (800 - sub_width) // 2
            draw.text((sub_x, 240), sub_text, fill=text_color, font=font_medium)
        else:
            # Fallback without font
            draw.text((250, 200), main_text, fill=text_color)
            draw.text((300, 240), sub_text, fill=text_color)
        
        # Convert to bytes
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        
        return base64.b64encode(img_bytes.read()).decode('utf-8')
    except Exception as e:
        print(f"Error creating AWS logo image: {str(e)}")
        # Return a simple base64 encoded 1x1 pixel image as fallback
        fallback_img = Image.new('RGB', (1, 1), color='white')
        fallback_bytes = io.BytesIO()
        fallback_img.save(fallback_bytes, format='PNG')
        fallback_bytes.seek(0)
        return base64.b64encode(fallback_bytes.read()).decode('utf-8')

@cloudformation_bp.route('/list-stacks', methods=['POST'])
def list_stacks():
    try:
        data = request.json
        region = data['region']
        access_key_id = data['accessKeyId']
        secret_access_key = data['secretAccessKey']
        query = data.get('query', '').lower()
        
        cf = create_aws_client('cloudformation', region, access_key_id, secret_access_key)
        
        stacks = []
        paginator = cf.get_paginator('list_stacks')
        
        for page in paginator.paginate(StackStatusFilter=[
            'CREATE_COMPLETE', 'UPDATE_COMPLETE', 'UPDATE_ROLLBACK_COMPLETE',
            'UPDATE_ROLLBACK_FAILED', 'CREATE_FAILED', 'UPDATE_FAILED', 'ROLLBACK_COMPLETE'
        ]):
            for stack in page.get('StackSummaries', []):
                stack_name = stack['StackName']
                if query in stack_name.lower():
                    stacks.append({
                        'name': stack_name,
                        'status': stack['StackStatus'],
                        'creationTime': stack['CreationTime'].isoformat()
                    })
        
        # Sort by name and limit results
        stacks.sort(key=lambda x: x['name'])
        return jsonify({'stacks': stacks[:20]})  # Limit to 20 results
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@cloudformation_bp.route('/generate-implementation-plan', methods=['POST'])
def generate_implementation_plan():
    try:
        data = request.json
        crq_plan = data['crqPlan']
        environment = data['environment']
        account_id = data['accountId']
        stack_name = data['stackName']
        developer = data['developer']
        images = data.get('images', [])
        work_info_type = data.get('workInfoType', '')
        auto_generate = data.get('autoGenerate', False)
        
        # Backward compatibility: single image
        if not images and 'imageData' in data:
            images = [{'data': data['imageData'], 'name': data.get('imageName', 'screenshot.png')}]
        
        # For Test Plan auto-generation, enhance the image if it's auto-generated
        if work_info_type == '16000' and auto_generate and images:
            # The frontend already created the image, but we can enhance it here if needed
            # Keep the existing image from frontend
            print(f"Test Plan auto-generation: Using frontend-generated AWS logo image")
        
        doc_heading = data.get('docHeading', 'Implementation Plan')
        doc_name = data.get('docName', 'implementPlan.docx')
        
        print(f"Generating {doc_name}: {crq_plan}, Stack: {stack_name}")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(doc_heading, 0)
        title.alignment = 1  # Center
        
        # Add fields based on work info type
        if work_info_type == '23000':  # Post Implementation Review
            doc.add_paragraph(f'{doc_heading}: {crq_plan}', style='Heading 2')
            doc.add_paragraph(f'Environment: {environment}')
            doc.add_paragraph(f'Risk Assessment: LOW')
            doc.add_paragraph(f'AWS Account: {account_id}')
            doc.add_paragraph(f'Stack Name: {stack_name}')
            doc.add_paragraph(f'Change Description: {stack_name}')
            doc.add_paragraph(f'Implementation Date: {datetime.now().strftime("%Y/%m/%d %H:%M:%S")}')
            doc.add_paragraph(f'Developer: {developer}')
            doc.add_paragraph(f'Deployment Class: Normal')
            doc.add_paragraph(f'Deployment Impact: Minor/Localised')
            doc.add_paragraph(f'Change Priority: LOW')
            doc.add_paragraph(f'Type of Change: Enhancement [always on/always secure]')
            doc.add_paragraph(f'Backend: AWS Cloud-Cape Town region')
            doc.add_paragraph(f'UI: AWS Cloud-Cape Town region')
        elif work_info_type == '11000':  # Backout Plan
            doc.add_paragraph(f'{doc_heading}: {crq_plan}', style='Heading 2')
            doc.add_paragraph(f'Environment: {environment}')
            doc.add_paragraph(f'Backout Risk Assessment: LOW')
            doc.add_paragraph(f'AWS Account: {account_id}')
            doc.add_paragraph(f'Backout Stack Name: {stack_name}')
            doc.add_paragraph(f'Backout Change Description: {stack_name}')
            doc.add_paragraph(f'Backout Implementation Date: {datetime.now().strftime("%Y/%m/%d %H:%M:%S")}')
            doc.add_paragraph(f'Developer: {developer}')
            doc.add_paragraph(f'Backout Class: Normal')
            doc.add_paragraph(f'Backout Impact: Minor/Localised')
            doc.add_paragraph(f'Backout Priority: LOW')
            doc.add_paragraph(f'Type of Backout Change: Enhancement [always on/always secure]')
            doc.add_paragraph(f'Backend: AWS Cloud-Cape Town region')
            doc.add_paragraph(f'UI: AWS Cloud-Cape Town region')
        elif work_info_type == '20000':  # Test Results
            doc.add_paragraph(f'AWS Account: {account_id}')
            doc.add_paragraph(f'Stack Name: {stack_name}')
            doc.add_paragraph(f'Test Results Description: {stack_name}')
            doc.add_paragraph(f'Test Results Date: {datetime.now().strftime("%Y/%m/%d %H:%M:%S")}')
            doc.add_paragraph(f'Developer: {developer}')
        elif work_info_type == '16000':  # Test Plan
            # Specific content for Test Plan as requested - only these fields
            doc.add_heading('Test Plans', level=1)
            doc.add_paragraph(f'Stack Name: {stack_name}')
            doc.add_paragraph(f'Test Plan Description: {stack_name}')
            doc.add_paragraph(f'Developer: {developer}')
            doc.add_paragraph(f'Test Plan: Not Applicable for this change')
            doc.add_paragraph(f'Date: {datetime.now().strftime("%Y/%m/%d %H:%M:%S")}')
            
            # Skip the default content section for Test Plan
            doc.add_paragraph()
            doc.add_heading('Attachment', level=2)
            for img in images:
                image_bytes = base64.b64decode(img['data'])
                image_stream = io.BytesIO(image_bytes)
                doc.add_picture(image_stream, width=Inches(5))
                doc.add_paragraph()
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            # Return as base64
            doc_base64 = base64.b64encode(doc_bytes.read()).decode('utf-8')
            
            print(f"Test Plan document generated successfully, size: {len(doc_base64)} bytes")
            
            return jsonify({'documentData': doc_base64})
        elif work_info_type == '22000':  # Cancellation Information
            doc.add_paragraph(f'{doc_heading}: {crq_plan}', style='Heading 2')
            doc.add_paragraph(f'Environment: {environment}')
            doc.add_paragraph(f'Risk Assessment: LOW')
            doc.add_paragraph(f'AWS Account: {account_id}')
            doc.add_paragraph(f'Cancelled Stack Name: {stack_name}')
            doc.add_paragraph(f'Cancelled Change Description: {stack_name}')
            doc.add_paragraph(f'Cancelled Date: {datetime.now().strftime("%Y/%m/%d %H:%M:%S")}')
            doc.add_paragraph(f'Developer: {developer}')
            doc.add_paragraph(f'Cancelled Deployment Class: Normal')
            doc.add_paragraph(f'Cancelled Deployment Impact: Minor/Localised')
            doc.add_paragraph(f'Change Priority: LOW')
            doc.add_paragraph(f'Type of Change: Enhancement [always on/always secure]')
            doc.add_paragraph(f'Backend: AWS Cloud-Cape Town region')
            doc.add_paragraph(f'UI: AWS Cloud-Cape Town region')
        elif work_info_type == '14000':  # Install Plan
            doc.add_paragraph(f'Installation Plan: {crq_plan}', style='Heading 2')
            doc.add_paragraph(f'Environment: {environment}')
            doc.add_paragraph(f'AWS Account: {account_id}')
            doc.add_paragraph(f'Stack Name: {stack_name}')
            doc.add_paragraph(f'Change Description: {stack_name}')
            doc.add_paragraph(f'Change Plan Date: {datetime.now().strftime("%Y/%m/%d %H:%M:%S")}')
            doc.add_paragraph(f'Developer: {developer}')
            doc.add_paragraph(f'Deployment Type: CloudFormation Stack Deployment')
            doc.add_paragraph(f'Deployment Steps: Automated via AWS CloudFormation')
        else:  # Default content for other work types
            doc.add_paragraph(f'{doc_heading}: {crq_plan}', style='Heading 2')
            doc.add_paragraph(f'Environment: {environment}')
            doc.add_paragraph(f'Risk Assessment: LOW')
            doc.add_paragraph(f'AWS Account: {account_id}')
            doc.add_paragraph(f'Stack Name: {stack_name}')
            doc.add_paragraph(f'Change Description: {stack_name}')
            doc.add_paragraph(f'Implementation Date: {datetime.now().strftime("%Y/%m/%d %H:%M:%S")}')
            doc.add_paragraph(f'Developer: {developer}')
            doc.add_paragraph(f'Deployment Class: Normal')
            doc.add_paragraph(f'Deployment Impact: Minor/Localised')
            doc.add_paragraph(f'Change Priority: LOW')
            doc.add_paragraph(f'Type of Change: Enhancement [always on/always secure]')
            doc.add_paragraph(f'Backend: AWS Cloud-Cape Town region')
            doc.add_paragraph(f'UI: AWS Cloud-Cape Town region')
        
        # Add images at the end
        doc.add_paragraph()
        doc.add_heading('Attachment', level=2)
        for img in images:
            image_bytes = base64.b64decode(img['data'])
            image_stream = io.BytesIO(image_bytes)
            doc.add_picture(image_stream, width=Inches(5))
            doc.add_paragraph()
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        # Return as base64
        doc_base64 = base64.b64encode(doc_bytes.read()).decode('utf-8')
        
        print(f"Document generated successfully, size: {len(doc_base64)} bytes")
        
        return jsonify({'documentData': doc_base64})
    except Exception as e:
        print(f"Error generating document: {str(e)}")
        return jsonify({'error': str(e)}), 500